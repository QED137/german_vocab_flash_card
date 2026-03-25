from docx import Document
from pathlib import Path
import re

import schemas


SEPARATORS = [" – ", " - ", ": ", " — "]
LESSON_PATTERN = re.compile(r"^lektion\b", re.IGNORECASE)
DETAIL_LABELS = {
    "meaning (en)": "meaning",
    "example (de)": "example_sentence",
    "translation (en)": "translation",
}


def _clean_text(value: str) -> str:
    return " ".join(value.split()).strip()


def _is_heading(line: str) -> bool:
    normalized = _clean_text(line).lower()
    return bool(LESSON_PATTERN.match(normalized))


def _normalize_lesson_name(value: str) -> str | None:
    cleaned = _clean_text(value).strip("-_: ")
    return cleaned or None


def extract_docx_lesson(file_path: str) -> str | None:
    document = Document(file_path)

    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if _is_heading(text):
            return _normalize_lesson_name(text)

    return _normalize_lesson_name(Path(file_path).stem.replace("_", " "))


def _split_line(line: str) -> tuple[str, str] | None:
    cleaned = _clean_text(line)
    for separator in SEPARATORS:
        if separator in cleaned:
            left, right = cleaned.split(separator, 1)
            left = left.strip("-• \t")
            right = right.strip()
            if left and right:
                return left, right
    return None


def _extract_labeled_details(text: str) -> dict[str, str]:
    details: dict[str, str] = {}
    ordered_values: list[tuple[str, str]] = []

    for raw_line in text.splitlines():
        line = _clean_text(raw_line)
        if not line or ":" not in line:
            if line:
                ordered_values.append(("", line))
            continue

        label, value = line.split(":", 1)
        normalized_label = label.strip().lower()
        key = DETAIL_LABELS.get(normalized_label)
        cleaned_value = value.strip()
        if key and cleaned_value:
            details[key] = cleaned_value
            continue

        if normalized_label in {"en", "de"} and cleaned_value:
            ordered_values.append((normalized_label, cleaned_value))
            continue

        if cleaned_value:
            ordered_values.append(("", cleaned_value))

    if details.get("meaning"):
        return details

    values = [value for _, value in ordered_values if value]
    labels = [label for label, value in ordered_values if value]

    if labels[:3] == ["en", "de", "en"] and len(values) >= 1:
        details["meaning"] = values[0]
        if len(values) >= 2:
            details["example_sentence"] = values[1]
        if len(values) >= 3:
            details["translation"] = values[2]
        return details

    if values:
        details["meaning"] = values[0]
    if len(values) >= 2:
        details["example_sentence"] = values[1]
    if len(values) >= 3:
        details["translation"] = values[2]

    return details


def _parse_table_rows(document: Document, lesson: str | None) -> list[schemas.WordCreate]:
    parsed: list[schemas.WordCreate] = []

    for table in document.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            word = _clean_text(row.cells[0].text)
            raw_details = row.cells[1].text
            details = _extract_labeled_details(raw_details)
            meaning = details.get("meaning")

            if not word or not meaning:
                continue

            translation = details.get("translation")
            notes = f"Translation (EN): {translation}" if translation else None

            parsed.append(
                schemas.WordCreate(
                    english_word=word,
                    lesson=lesson,
                    meaning=meaning,
                    example_sentence=details.get("example_sentence"),
                    notes=notes,
                )
            )

    return parsed


def parse_docx_words(file_path: str) -> tuple[str | None, list[schemas.WordCreate]]:
    document = Document(file_path)
    lesson = extract_docx_lesson(file_path)
    parsed = _parse_table_rows(document, lesson)

    if parsed:
        return lesson, parsed

    parsed = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or _is_heading(text):
            continue

        split_result = _split_line(text)
        if not split_result:
            continue

        word, meaning = split_result
        parsed.append(
            schemas.WordCreate(
                english_word=word,
                lesson=lesson,
                meaning=meaning,
            )
        )

    return lesson, parsed
